from pathlib import Path

import fitz


def pdf_to_docx(input_pdf: Path, output_docx: Path) -> dict:
    """Convert PDF to DOCX using text extraction and python-docx."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx belum terinstall. Jalankan: pip install python-docx")

    try:
        doc_pdf = fitz.open(input_pdf)
    except Exception as exc:
        raise ValueError("PDF_CORRUPTED") from exc

    if doc_pdf.is_encrypted:
        doc_pdf.close()
        raise PermissionError("PDF_ENCRYPTED")

    total_pages = doc_pdf.page_count
    if total_pages <= 0:
        doc_pdf.close()
        raise ValueError("PDF_CORRUPTED")

    document = Document()

    style = document.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    for page_index in range(total_pages):
        page = doc_pdf.load_page(page_index)

        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]

        for block in blocks:
            if block["type"] == 0:
                for line in block["lines"]:
                    line_text = ""
                    for span in line["spans"]:
                        line_text += span["text"]
                    line_text = line_text.rstrip()
                    if line_text:
                        para = document.add_paragraph(line_text)
                        span_info = line["spans"][0] if line["spans"] else None
                        if span_info and span_info["size"] > 14:
                            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            for run in para.runs:
                                run.bold = True
                                run.font.size = Pt(min(int(span_info["size"]), 28))
            elif block["type"] == 1:
                try:
                    img_data = block.get("image")
                    if img_data:
                        img_path = output_docx.parent / f"_img_{page_index}_{block['number']}.png"
                        img_path.write_bytes(img_data)
                        document.add_picture(str(img_path), width=Inches(5.5))
                        img_path.unlink(missing_ok=True)
                except Exception:
                    pass

        if page_index < total_pages - 1:
            document.add_page_break()

    doc_pdf.close()
    document.save(str(output_docx))

    return {
        "totalPages": total_pages,
        "originalSize": input_pdf.stat().st_size,
        "outputSize": output_docx.stat().st_size,
    }
